"""
Module tools docx of MCQMaker

It contains all functions to manipulate docx documents.

Functions
---------
"""

###############
### Imports ###
###############

import os
# Import deepcopy to dupplicate paragraphs
from copy import deepcopy

# Import docx to create word documents
import docx
from docx import Document

from mcq_maker_tools.tools import (
    PATH_TEMPLATE_FOLDER,
    filter_hidden_files
)

#################
### Functions ###
#################

def get_list_templates():
    """
    Return the list of names of the templates stored in the template folder.
    """
    template_files_list = os.listdir(PATH_TEMPLATE_FOLDER)
    cleaned_template_files_list = filter_hidden_files(
        template_files_list, ".docx")
    res = [e.replace(".docx", "") for e in cleaned_template_files_list]
    return res

### Error correction ###


def extract_style(run: docx.text.run.Run):
    """
    Extract the style of a run and returns it in a tuple.

    Parameters
    ----------
    run : docx.text.run.Run
        The run where to extract the style

    Returns
    -------
    tuple
    """
    return (run.style, run.bold, run.italic, run.underline)


# Fonction corrigeant les problèmes de runs dans un paragraphe
def check_error_in_paragraph(paragraph: docx.text.paragraph.Paragraph):
    if len(paragraph.runs) > 0 and not("Graphic" in paragraph._p.xml):
        sto = extract_style(paragraph.runs[0])
        j_sto = 0
        l_text = [""]
        l_j = []
        for j in range(len(paragraph.runs)):
            cur_style = extract_style(paragraph.runs[j])
            if sto == cur_style:
                l_text[-1] += paragraph.runs[j].text
                if j != 0:
                    l_j.append(j)
                    paragraph.runs[j_sto].text += paragraph.runs[j].text
            else:
                sto = cur_style
                j_sto = j
                l_text.append(paragraph.runs[j].text)
        for j in l_j:
            paragraph.runs[j].clear()


# Fonction corrigeant les problèmes de runs dans le document Word (si on modifie seulement la moitié d'un mot dans le Word, ça fait n'importe quoi par exemple)
def check_error_in_codes(document):

    # Corriger le problème pour les paragraphes du document
    for i_para in range(len(document.paragraphs)):
        check_error_in_paragraph(document.paragraphs[i_para])

    # Corriger le problème pour les tableaux du document
    for i_table in range(len(document.tables)):
        for i_row in range(len(document.tables[i_table].rows)):
            for i_cell in range(len(document.tables[i_table].rows[i_row].cells)):
                for i_para in range(len(document.tables[i_table].rows[i_row].cells[i_cell].paragraphs)):
                    check_error_in_paragraph(
                        document.tables[i_table].rows[i_row].cells[i_cell].paragraphs[i_para])


### Deletion ###

# Fonction supprimant les paragraphes d'indications avec les ###
def delete_para_indications(paragraph):
    if "###" in paragraph.text:
        paragraph._element.getparent().remove(paragraph._element)


# Fonction supprimant tous les paragraphes d'indications avec les ###
def delete_indications(document: Document):
    for paragraph in document.paragraphs:
        delete_para_indications(paragraph)
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    delete_para_indications(paragraph)

### Find ###

# Fonction retournant le dictionnaire qui indique le début de chaque paragraphe de coupure
def detect_cut_para(document):
    dict_begin_para = {}
    for index_para in range(len(document.paragraphs)):
        if "###" in document.paragraphs[index_para].text and "START" in document.paragraphs[index_para].text:
            dict_begin_para[document.paragraphs[index_para].text] = [
                document.paragraphs[index_para], index_para]
    return dict_begin_para


# Fonction trouvant un texte dans un paragraphe et qui renvoie le paragraphe en question et son indice
def find_paragraph(document, text_to_find):
    for index_para in range(len(document.paragraphs)):
        if text_to_find in document.paragraphs[index_para].text:
            return document.paragraphs[index_para], index_para
    return None, 0


# Fonction détectant les mots-clés à remplacer dans un paragraphe et les stocke dans une liste
def find_keyword_paragraph(paragraph):
    list_keywords = []
    text_paragraph = paragraph.text
    keyword = ""
    keyword_detection = False
    for character in text_paragraph:
        if character == '{':
            keyword_detection = True
        if character == '}':
            keyword += character
            list_keywords.append(keyword)
            keyword_detection = False
            keyword = ""
        if keyword_detection:
            keyword += character
    return list_keywords


# Fonction détectant les mots-clés à remplacer dans un paragraphe ou une table et les renvoie sous forme d'ensemble
def find_keywords(element, type):
    if type == "paragraph":
        list_keywords = find_keyword_paragraph(element)
    elif type == "table":
        list_keywords = []
        for row in element.rows:
            for paragraph in row.cells:
                list_keywords += find_keyword_paragraph(paragraph)
    set_keywords = set(list_keywords)
    return set_keywords

### Replacements ###

def replace_text_paragraph(paragraph, new_text):
    paragraph.runs[0].text = new_text
    for i in range(1, len(paragraph.runs)):
        paragraph.runs[i].clear()


def reformat_replacement(replacement):
    replacement = replacement.replace("  ", "")
    replacement = replacement.replace("\t", "")
    l_para = replacement.split("\n")
    l_res = []
    for i in range(len(l_para)):
        if len(l_para[i]) > 0:
            if l_para[i][0] == "•" or l_para[i][0] == "-":
                l_res.append([l_para[i], "list"])
            else:
                l_res.append([l_para[i], "normal"])
        else:
            l_res.append([" ", "normal"])
    return l_res


# Fonction remplaçant le mot-clé dans un paragraphe
def replace_in_paragraph(paragraph, keyword, replacement, debug_mode=False):
    if replacement == "nan":
        replacement = "aucun"
    if keyword == paragraph.text:
        # Dans ce cas-ci, on va appliquer un reformattage de replacement pour l'écrire sur plusieurs paragraphes
        replacement = reformat_replacement(replacement)
        for id in range(len(replacement) - 1, -1, -1):
            if id == 0:
                replace_text_paragraph(paragraph, replacement[0][0])
            else:
                new_p = deepcopy(paragraph)
                replace_text_paragraph(new_p, replacement[id][0])
                paragraph._p.addnext(new_p._p)

    elif keyword in paragraph.text:
        for j in range(len(paragraph.runs)):
            if keyword in paragraph.runs[j].text:
                # Majuscule automatique si le caractère précédent est une tabulation
                if paragraph.runs[j - 1].text != "\t":
                    paragraph.runs[j].text = paragraph.runs[j].text.replace(
                        keyword, replacement)
                else:
                    paragraph.runs[j].text = paragraph.runs[j].text.replace(
                        keyword, replacement[0].upper() + replacement[1:])

def replace_in_doc(document, keyword, replacement, debug_mode=False):
    """
    Replace a keyword in the whole document.

    Parameters
    ----------
    keyword : str
        String containg the keyword to replace

    replacement : str
        String containing the replacement.

    Returns
    -------
    None
    """
    for paragraph in document.paragraphs:
        replace_in_paragraph(paragraph, keyword, replacement, debug_mode)

# Fonction remplaçant dans le tableau de référence le mot-clé souhaité par la valeur souhaitée, tout en conservant le style graphique du tableau
def replace_in_table(ref_table, keyword, replacement):
    for i_row in range(len(ref_table.rows)):
        for i_cell in range(len(ref_table.rows[i_row].cells)):
            for i_para in range(len(ref_table.rows[i_row].cells[i_cell].paragraphs)):
                replace_in_paragraph(
                    ref_table.rows[i_row].cells[i_cell].paragraphs[i_para], keyword, replacement)
