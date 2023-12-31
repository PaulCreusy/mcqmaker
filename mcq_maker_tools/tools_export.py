
###############
### Imports ###
###############

### Python imports ###

import sys
import os
from copy import deepcopy

sys.path.append(".")

import random
from lxml import etree
import copy
from docx import Document

### Module imports ###

from mcq_maker_tools.tools import (
    SETTINGS,
    PATH_TEMPLATE_FOLDER,
    DICT_LANGUAGE,
    convert_int_to_letter,
    remove_three_dots
)
from mcq_maker_tools.tools_class import (
    load_class,
    save_class
)
from mcq_maker_tools.tools_database import (
    load_database
)
from mcq_maker_tools.tools_docx import (
    replace_in_doc,
    find_paragraph,
    delete_indications
)


#################
### Functions ###
#################


### QCM functions ###

def generate_MCQ(config, class_content):
    """
    Generate the QCM data to then export it in the selected format.

    Parameters
    ----------
    config : dict
        Configuration dictionnary under the following form :
        {
            "QCM_name": str,
            "questions":
                [
                    {"folder_name": str, "file_name": str, "nb_questions": int},
                ],
            "template": str,
            "mix_all_questions": bool,
            "mix_among_databases": bool
        }

    Returns
    -------
    dict
        QCM data under the following form :
        {
            "QCM_name": str,
            "questions": [
                {"question": str, "options": list, "answer": int}
            ]
        }
    """

    class_content = copy.copy(class_content)
    config = copy.copy(config)

    # Extract information from the config dict
    mix_all_questions = config["mix_all_questions"]
    mix_among_databases = config["mix_among_databases"]
    instructions = config["questions"]

    # Initialise data structures
    questions_sublists = []
    questions = []
    QCM_data = {}
    QCM_data["QCM_name"] = config["QCM_name"]
    QCM_data["template"] = config["template"]

    # Error boolean
    error_detected = False
    l_error_files = []

    # Build sublist of questions
    for i in range(len(instructions)):

        # Load the data
        current_dict = instructions[i]
        folder = current_dict["folder_name"].replace("\n", " ")
        file = current_dict["file_name"].replace("\n", " ")
        nb_questions = current_dict["nb_questions"]
        database_questions = load_database(file, folder)

        # Remove already selected questions
        if (folder, file) in class_content:
            used_questions_list = class_content[(
                folder, file)]["list_questions_used"]
            for i in range(len(database_questions) - 1, -1, -1):
                el = database_questions[i]
                if el["id"] in used_questions_list:
                    database_questions.pop(i)

        # Mix if needed
        if mix_among_databases:
            selected_questions_id = random.sample(
                [j for j in range(len(database_questions))], nb_questions)
        else:
            selected_questions_id = [i for i in range(
                len(database_questions))][:nb_questions]

        # Convert the selected questions id in a format that can be save for the class
        selected_id_to_save = [database_questions[i]["id"]
                               for i in selected_questions_id]

        # Insert the selected questions in the class content
        if (folder, file) in class_content:
            class_content[(folder, file)
                          ]["list_questions_used"] += selected_id_to_save
            class_content[(folder, file)
                          ]["used_questions"] += len(selected_id_to_save)
        else:
            class_content[(folder, file)] = {
                "list_questions_used": selected_id_to_save, "used_questions": len(selected_id_to_save)}

        # Extract the questions using the id
        selected_questions = [database_questions[j]
                              for j in selected_questions_id]

        # Insert inside the list
        questions_sublists.append(selected_questions)

    if error_detected:
        return None, l_error_files

    # Merge all sublists into a single one
    for sublist in questions_sublists:
        questions += sublist

    # Mix all questions if the option is selected
    if mix_all_questions:
        random.shuffle(questions)

    # Insert data inside the dict
    QCM_data["questions"] = questions

    return QCM_data, class_content


def create_folder_MCQ(QCM_data):
    QCM_name = QCM_data["QCM_name"]
    folder_path = SETTINGS["path_export"] + QCM_name
    new_folder_path = folder_path
    i = 1
    while os.path.exists(new_folder_path):
        new_folder_path = folder_path + "_" + str(i)
        i += 1

    os.mkdir(new_folder_path)

    return new_folder_path + "/"


def export_MCQ_txt(QCM_data, folder_path, progress_bar):
    """
    Export the QCM and its solution in a .txt file.

    Parameters
    ----------
    QCM_data : dict
        Data to generate the QCM under the following form :
        {
            "QCM_name": str,
            "questions": [
                {"question": str, "options": list, "answer": int}
            ]
        }

    progress_bar : ProgressBar
        Kivy progress bar to update it on the interface.

    Returns
    -------
    None
    """

    # Build the path of the files
    QCM_path = folder_path + QCM_data["QCM_name"] + ".txt"
    solution_path = folder_path + QCM_data["QCM_name"] + "_solution.txt"

    # Open the files
    QCM_file = open(QCM_path, "w", encoding="utf-8")
    solution_file = open(solution_path, "w", encoding="utf-8")

    # Extract information from the dictionnary
    QCM_name = QCM_data["QCM_name"]
    questions = QCM_data["questions"]

    # Write the first lines
    solution_file.write(QCM_name + "\n\n")
    solution_file.write("Id de la question\tBonne reponse\n")
    QCM_file.write(QCM_name + "\n\n")
    QCM_file.write("Nom : " + " " * 20 + "Prénom : " + " " * 20 + "\n")
    QCM_file.write("Classe :\n\n")

    progress_bar.value = 3
    number_questions = len(questions)

    # Scan the list of questions
    for i in range(number_questions):

        # Extract the data
        question_dict = questions[i]
        question = question_dict["question"]
        options = question_dict["options"]
        answer = question_dict["answer"]

        # Write the current question
        QCM_file.write(f"{str(i+1)}. {question}\n")
        for j in range(len(options)):
            QCM_file.write(f"\t{convert_int_to_letter(j)}) {options[j]}")
        QCM_file.write("\n\n")

        # Write the solution
        solution_file.write(f"{i}\t{convert_int_to_letter(answer)}\n")

        # Update the value of the progress bar
        progress_bar.value += 17 / number_questions


def export_MCQ_docx(QCM_data, folder_path, progress_bar):

    # Create the file object with the selected template
    QCM_file = Document(PATH_TEMPLATE_FOLDER + QCM_data["template"] + ".docx")

    # Replace the name of the MCQ
    replace_in_doc(QCM_file, "{NAME_MCQ}", QCM_data["QCM_name"])

    # Detect the id of the paragraph where the list of questions starts
    begin_para, begin_para_idx = find_paragraph(
        QCM_file, "### LIST_QUESTIONS_START ###")
    end_para, end_para_idx = find_paragraph(
        QCM_file, "### LIST_QUESTIONS_END ###")

    # Store in a list all paragraphs composing a question to dupplicate them
    para_list = []
    for i in range(begin_para_idx + 1, end_para_idx):
        para_list.append(QCM_file.paragraphs[i])

    copy_para_list = deepcopy(para_list)

    # Raise error if no paragraph is found
    if begin_para is None or end_para is None:
        raise ValueError(
            "The selected template does not include a list of question area.")

    progress_bar.value = 23

    nb_questions = len(QCM_data["questions"])

    # Add the questions to the document
    for (i, question_dict) in enumerate(QCM_data["questions"]):
        if i > 0:
            # Do a copy of the list to create the new paragraphs
            para_list = deepcopy(copy_para_list)
            for (j, para) in enumerate(para_list):
                new_para = para._p
                QCM_file.paragraphs[end_para_idx]._p.addnext(new_para)
                end_para_idx += 1

        # Replace the id and the question
        replace_in_doc(QCM_file, "{ID_QUESTION}", str(i + 1))
        replace_in_doc(QCM_file, "{QUESTION}", question_dict["question"])

        # Locate the paragraph containing the list of options
        options = question_dict["options"]
        options_para, end_para_idx = find_paragraph(QCM_file, "{LIST_OPTIONS}")
        copy_options_para = deepcopy(options_para)

        # Add the first option
        replace_in_doc(QCM_file, "{LIST_OPTIONS}",
                       convert_int_to_letter(0) + ". " + options[0])

        # Dupplicate the paragraph to add the answers
        nb_answers = len(options)
        for j in range(1, nb_answers):

            new_copy_options_para = deepcopy(copy_options_para)

            # Dupplicate the list options
            new_para = copy_options_para._p
            QCM_file.paragraphs[end_para_idx]._p.addnext(new_para)

            # Replace it with the current option
            replace_in_doc(QCM_file, "{LIST_OPTIONS}",
                           convert_int_to_letter(j) + ". " + options[j])

            copy_options_para = new_copy_options_para

            end_para_idx += 1

        progress_bar.value += 15 / nb_questions

    delete_indications(QCM_file)

    QCM_file.save(folder_path + QCM_data["QCM_name"] + ".docx")

    progress_bar.value = 40

def export_MCQ_H5P_text_single_choice(QCM_data, folder_path, progress_bar):
    """
    Export the QCM and its solution in a .h5p file for sinle choice H5P integration.

    Parameters
    ----------
    QCM_data : dict
        Data to generate the QCM under the following form :
        {
            "QCM_name": str,
            "questions": [
                {"question": str, "options": list, "answer": int}
            ]
        }

    progress_bar : ProgressBar
        Kivy progress bar to update it on the interface.

    Returns
    -------
    None
    """

    # Define the path of the file
    file_path = folder_path + \
        QCM_data["QCM_name"] + "_H5P_text_single_choice.txt"

    nb_questions = len(QCM_data["questions"])

    # Open the file to write it
    with open(file_path, "w", encoding="utf-8") as file:
        for question_dict in QCM_data["questions"]:

            # Put the answer in first position
            answer_id = question_dict["answer"]
            crossed_list = [(i != answer_id, e)
                            for (i, e) in enumerate(question_dict["options"])]
            crossed_list = sorted(crossed_list)
            options_list = [e[1] for e in crossed_list]

            # Ecriture dans le fichier
            file.write(question_dict["question"] + "\n")
            for option in options_list:
                file.write(option + "\n")
            file.write("\n")

            progress_bar.value += 20 / nb_questions


def export_MCQ_H5P_text_fill_blanks(QCM_data, folder_path, progress_bar):
    """
    Export the QCM and its solution in a .h5p file for sinle choice H5P integration.

    Parameters
    ----------
    QCM_data : dict
        Data to generate the QCM under the following form :
        {
            "QCM_name": str,
            "questions": [
                {"question": str, "options": list, "answer": int}
            ]
        }

    progress_bar : ProgressBar
        Kivy progress bar to update it on the interface.

    Returns
    -------
    None
    """

    # Define the path of the file
    file_path = folder_path + \
        QCM_data["QCM_name"] + "_H5P_text_fill_in_the_blanks.txt"

    nb_questions = len(QCM_data["questions"])

    # Open the file to write it
    with open(file_path, "w", encoding="utf-8") as file:
        for question_dict in QCM_data["questions"]:

            # Split the question
            question = question_dict["question"]
            question = remove_three_dots(question)
            split_question = question.split("...", 1)

            # Create the line
            answer = question_dict["options"][question_dict["answer"]]
            line = split_question[0] + " *" + answer + "* " + split_question[1]

            # Ecriture dans le fichier
            file.write(line + "\n")

            progress_bar.value += 20 / nb_questions


def export_MCQ_moodle(QCM_data, folder_path, progress_bar):
    """
    Export the QCM and its solution in a .xml file for moodle.

    Parameters
    ----------
    QCM_data : dict
        Data to generate the QCM under the following form :
        {
            "QCM_name": str,
            "questions": [
                {"question": str, "options": list, "answer": int}
            ]
        }

    progress_bar : ProgressBar
        Kivy progress bar to update it on the interface.

    Returns
    -------
    None
    """

    # Build the path of the file
    QCM_path = folder_path + QCM_data["QCM_name"] + ".xml"

    ### Build the xml tree ###

    # Introduction

    QCM_tree = etree.Element("quiz")

    QCM_intro = etree.SubElement(QCM_tree, "question")
    QCM_intro.set("type", "category")

    QCM_intro_cat = etree.SubElement(QCM_intro, "category")
    QCM_intro_cat_txt = etree.SubElement(QCM_intro_cat, "text")
    QCM_intro_cat_txt.text = QCM_data["QCM_name"]

    QCM_intro_info = etree.SubElement(QCM_intro, "info")
    QCM_intro_info.set("format", "moodle_auto_format")

    QCM_intro_info_txt = etree.SubElement(QCM_intro_info, "text")
    QCM_intro_info_txt.text = "QCM"

    QCM_intro_id = etree.SubElement(QCM_intro, "idnumber")

    progress_bar.value = 63

    # Questions

    # Extract the data
    questions = QCM_data["questions"]

    progress_bar.value = 83
    nb_questions = len(QCM_data["questions"])

    for i in range(len(questions)):
        question_dict = questions[i]

        # Create the tree of the question
        question = etree.SubElement(QCM_tree, "question")
        question.set("type", "multichoice")

        # Add the instruction
        question_name = etree.SubElement(question, "name")
        question_name_txt = etree.SubElement(question_name, "text")
        question_name_txt.text = "Choose the correct answer"

        # Add the question text
        question_text = etree.SubElement(question, "questiontext")
        question_text.set("format", "html")
        question_text_txt = etree.SubElement(question_text, "text")
        question_text_txt.text = "<![CDATA[<p>" + \
            question_dict["question"] + "<br></p>]]>"

        # Set the options of the question
        question_fb = etree.SubElement(question, "generalfeedback")
        question_fb.set("format", "html")
        question_fb_txt = etree.SubElement(question_fb, "text")
        question_fb_txt.text = ""

        question_dg = etree.SubElement(question, "defaultgrade")
        question_dg.text = "1.0"

        question_py = etree.SubElement(question, "penalty")
        question_py.text = "0.0"

        question_hi = etree.SubElement(question, "hidden")
        question_hi.text = "0"

        question_id = etree.SubElement(question, "idnumber")

        question_single = etree.SubElement(question, "single")
        question_single.text = "true"

        question_sa = etree.SubElement(question, "shuffleanswers")
        question_sa.text = "true"

        question_an = etree.SubElement(question, "answernumbering")
        question_an.text = "abc"

        question_ssi = etree.SubElement(question, "showstandardinstruction")
        question_ssi.text = "1"

        question_cfb = etree.SubElement(question, "correctfeedback")
        question_cfb.set("format", "html")
        question_cfb_txt = etree.SubElement(question_cfb, "text")
        question_cfb_txt.text = "Your answer is correct."

        question_pcfb = etree.SubElement(question, "partiallycorrectfeedback")
        question_pcfb.set("format", "html")
        question_pcfb_txt = etree.SubElement(question_pcfb, "text")
        question_pcfb_txt.text = "Your answer is partially correct."

        question_ifb = etree.SubElement(question, "partiallycorrectfeedback")
        question_ifb.set("format", "html")
        question_ifb_txt = etree.SubElement(question_ifb, "text")
        question_ifb_txt.text = "Your answer is incorrect."

        question_snc = etree.SubElement(question, "shownumcorrect")

        for j in range(len(question_dict["options"])):
            option = question_dict["options"][j]
            if j == question_dict["answer"]:
                fraction = "100"
            else:
                fraction = "0"
            question_answer = etree.SubElement(question, "answer")
            question_answer.set("fraction", fraction)
            question_answer.set("format", "html")
            question_answer_txt = etree.SubElement(question_answer, "text")
            question_answer_txt.text = "<![CDATA[<p>" + \
                option + "<br></p>]]>"
            question_answer_fb = etree.SubElement(question_answer, "feedback")
            question_answer_fb.set("format", "html")
            question_answer_fb_txt = etree.SubElement(
                question_answer_fb, "text")

        # Update the value of the progress bar
        progress_bar.value += 17 / nb_questions

    # Open the files
    QCM_file = open(QCM_path, "w", encoding="utf-8")

    # Write the xml in the file
    QCM_file.write(
        etree.tostring(QCM_tree, encoding="utf-8", pretty_print=True).decode('utf-8').replace("&lt;", "<").replace("&gt;", ">"))


def launch_export_MCQ(config, class_name, dict_formats, progress_bar, close_button, label_popup, popup):
    """
    Export the QCM in txt, xml for moodle and docx if a template is choosen and save the data in the class.

    Parameters
    ----------
    config : dict
        Dictionnary of configuration.

    class_name : str
        Name of the selected class can also be None.

    dict_formats : dict
        Dictionary containing the formats wanted by the user.

    progress_bar : ProgressBar
        Kivy progress bar to update the interface.

    Returns
    -------
    None
    """

    # Load the data of the class
    if class_name is not None:
        class_content = load_class(class_name)
    else:
        class_content = {}

    # Create the data of the QCM
    QCM_data, class_content = generate_MCQ(config, class_content)

    if QCM_data is None:
        files_string = ""
        for e in class_content:
            files_string += e + "\n"
        label_popup.text = DICT_LANGUAGE["qcm"]["qcm_generation"]["label_error_popup"] + files_string
        popup.title = DICT_LANGUAGE["qcm"]["qcm_generation"]["title_error_popup"]
        return False

    # Create the folder
    folder_path = create_folder_MCQ(QCM_data)

    # Export it as txt
    if dict_formats["txt"]:
        export_MCQ_txt(QCM_data, folder_path, progress_bar)

    # Export it as docx if a template is choosen
    if dict_formats["docx"]:
        try:
            export_MCQ_docx(QCM_data, folder_path, progress_bar)
        except:
            label_popup.text = DICT_LANGUAGE["qcm"]["qcm_generation"]["error_export_docx"]
            popup.title = DICT_LANGUAGE["qcm"]["qcm_generation"]["title_error_popup"]
            return False

    # Export it in xml for moodle
    if dict_formats["xml"]:
        export_MCQ_moodle(QCM_data, folder_path, progress_bar)

    # Export it in single choice H5P
    if dict_formats["h5p"]:
        export_MCQ_H5P_text_single_choice(QCM_data, folder_path, progress_bar)

    # Export it in fill in blanks H5P
    if dict_formats["h5p"]:
        try:
            export_MCQ_H5P_text_fill_blanks(
                QCM_data, folder_path, progress_bar)
        except:
            label_popup.text = DICT_LANGUAGE["qcm"]["qcm_generation"]["error_no_ellipsis"]
            popup.title = DICT_LANGUAGE["qcm"]["qcm_generation"]["title_error_popup"]
            return False

    # Save the class data if one is choosen
    if class_name is not None and config["update_class"]:
        save_class(class_name, class_content)

    # Final update of the content of the popup
    progress_bar.value = 100
    close_button.disabled = False
    label_popup.text = DICT_LANGUAGE["qcm"]["qcm_generation"]["label_end_popup"]
    popup.title = DICT_LANGUAGE["qcm"]["qcm_generation"]["title_end_popup"]
    return True
